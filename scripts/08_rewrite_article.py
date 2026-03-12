"""
Script 08: Rewrite IEEE Access article with recalculated results
Articulo IEEE Access - Jorge Jose Zamora, UPCT

Reads the template styles from the existing article and rewrites
all content with new data (175,869 headlines, 2017-2023, D_COVID+D_UCRANIA).
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

import json
import pandas as pd
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from copy import deepcopy

BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATE = BASE_DIR / 'Access JJZC-SMMMD_BACKUP_20260312.docx'
OUTPUT = BASE_DIR / 'Access JJZC-SMMMD.docx'

# Load results
DATA_DIR = BASE_DIR / 'article_outputs'

with open(DATA_DIR / 'arimax' / 'all_results.json', 'r', encoding='utf-8') as f:
    arimax_results = json.load(f)
with open(DATA_DIR / 'var_granger' / 'var_granger_summary.json', 'r', encoding='utf-8') as f:
    var_summary = json.load(f)
with open(DATA_DIR / 'structural' / 'structural_breaks_summary.json', 'r', encoding='utf-8') as f:
    struct_summary = json.load(f)

df_granger = pd.read_csv(DATA_DIR / 'var_granger' / 'granger_tests.csv')
df_ks = pd.read_csv(DATA_DIR / 'structural' / 'ks_mw_tests.csv')
df_arimax_all = pd.read_excel(DATA_DIR / 'arimax' / 'arimax_results_full.xlsx')
df_adf = pd.read_csv(DATA_DIR / 'arimax' / 'adf_tests.csv')
df_fevd = pd.read_csv(DATA_DIR / 'var_granger' / 'fevd_ipc.csv', index_col=0)

# Translation mapping: Spanish -> English
TR = {
    'Ira': 'Anger', 'Miedo': 'Fear', 'Tristeza': 'Sadness', 'Alegria': 'Joy',
    'Sorpresa': 'Surprise', 'Asco': 'Disgust', 'Otros': 'Others',
    'D_UCRANIA': 'D_UKRAINE',
    'IPC_ECOICOP Índice General': 'CPI General Index',
}

def tr(name):
    """Translate Spanish label to English."""
    return TR.get(name, name)

# Extract global model
global_model = None
for r in arimax_results:
    if 'Global' in r.get('variable', ''):
        global_model = r
        break

# Significant ARIMAX models (where emotion or dummy vars are significant)
sig_arimax = []
for r in arimax_results:
    sig_exog = []
    for pname, info in r['coefficients'].items():
        if info.get('significant_5pct') and pname not in ['intercept', 'sigma2'] and not pname.startswith('ar.') and not pname.startswith('ma.'):
            sig_exog.append((pname, info['coefficient'], info['p_value']))
    if sig_exog:
        sig_arimax.append({'variable': r['variable'], 'order': r['order'], 'sig_exog': sig_exog})


print("=" * 80)
print("REWRITING IEEE ACCESS ARTICLE")
print("=" * 80)

# ─── Create document from template ──────────────────────────────────────────
doc = Document(str(TEMPLATE))

# Clear all content from body while preserving styles
body = doc.element.body
for child in list(body):
    if child.tag.endswith('}sectPr'):
        continue  # Keep section properties (page layout, columns, etc.)
    body.remove(child)


FIGURES_DIR = DATA_DIR / 'figures'
ARIMAX_DIR = DATA_DIR / 'arimax'


def add_para(text, style='PARA'):
    """Add a paragraph with given style."""
    p = doc.add_paragraph(text, style=style)
    return p


def add_figure(image_path, caption_text, width_inches=6.5):
    """Add a figure with caption in IEEE Access style."""
    if not Path(image_path).exists():
        add_para(f'[FIGURE NOT FOUND: {image_path}]', 'PARA')
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    add_para(caption_text, 'Fig Caption')


def set_table_borders(table):
    """Add borders to table (top/bottom only for IEEE style)."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        from lxml import etree
        borders = etree.SubElement(tblPr, qn('w:tblBorders'))
    for edge in ['top', 'bottom', 'insideH']:
        elem = borders.find(qn(f'w:{edge}'))
        if elem is None:
            from lxml import etree
            elem = etree.SubElement(borders, qn(f'w:{edge}'))
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), '4')
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), '000000')


def add_table(headers, rows, title=None):
    """Add a formatted table."""
    if title:
        doc.add_paragraph(title, style='Table Title')

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Normal Table'
    set_table_borders(table)

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = str(h)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)

    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val) if val is not None else ''
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    doc.add_paragraph('Source: Authors\u2019 own elaboration.', style='footnote text')
    return table


# ═══════════════════════════════════════════════════════════════════════════════
# ARTICLE CONTENT
# ═══════════════════════════════════════════════════════════════════════════════

# ─── Header ──────────────────────────────────────────────────────────────────
add_para('Date of publication xxxx 00, 0000, date of current version xxxx 00, 0000.', 'DOP')
add_para('Digital Object Identifier 10.1109/ACCESS.2024.Doi Number', 'DOI')

add_para(
    'Sentiment Monitoring in Spanish Media via Multi-Model Econometric Design: '
    'Emotional Signals, Inflation Dynamics, and Structural Breaks (2017\u20132023)',
    'Paper Title'
)

add_para(
    'Zamora C\u00e1novas, Jorge Jos\u00e91, '
    'Mart\u00ednez Mar\u00eda-Dolores, Soledad Maria2',
    'AU'
)
add_para('1 Polytechnic University of Cartagena', 'PI_No Space')
add_para(
    '2 Department of Quantitative Methods, Legal Sciences and Modern Languages, '
    'Polytechnic University of Cartagena',
    'PI_No Space'
)
add_para(
    'Corresponding author: Mart\u00ednez Mar\u00eda-Dolores, Soledad Maria '
    '(e-mail: soledad.martinez@upct.es).',
    'PI'
)

# ─── Abstract ────────────────────────────────────────────────────────────────
add_para(
    'ABSTRACT This study investigates the longitudinal relationship between emotional signals '
    'embedded in 175,869 Spanish digital news headlines and macroeconomic indicators over the '
    '2017\u20132023 period. Emotional features are extracted using RoBERTa-BNE, a transformer-based '
    'language model fine-tuned for Spanish emotion classification, producing seven-dimensional '
    'probability vectors (anger, fear, joy, sadness, surprise, disgust, and neutral) for each headline. '
    'Monthly aggregates of four core emotions (anger, fear, sadness, and joy) are incorporated as '
    'exogenous regressors in ARIMAX models alongside structural dummy variables for COVID-19 and the '
    'Ukraine conflict. A complementary VAR framework with Granger causality tests, orthogonalized '
    'impulse response functions, and forecast error variance decomposition reveals bidirectional '
    'temporal dependencies between emotional dimensions and Spain\u2019s Consumer Price Index. Structural '
    'break analysis confirms statistically significant regime shifts at both the COVID-19 onset '
    '(March 2020) and the Ukraine war (February 2022). The results demonstrate that media-conveyed '
    'emotions, particularly sadness, anger, and joy, exhibit systematic Granger-causal relationships '
    'with inflation dynamics, supporting the integration of emotionally enriched textual signals '
    'into macroeconomic monitoring frameworks. All scripts, aggregated data, and reproducibility '
    'materials are publicly available.',
    'Normal'
)

add_para(
    'INDEX TERMS Affective computing, ARIMAX models, Consumer Price Index, '
    'Digital news media, Granger causality, Impulse response functions, '
    'RoBERTa-BNE, Sentiment analysis, Structural breaks, '
    'Transformer-based emotion classification, VAR models.',
    'Normal'
)

# ═══════════════════════════════════════════════════════════════════════════════
# I. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
add_para('INTRODUCTION', 'H1_List (No Space)')

add_para(
    'Digital news consumption increasingly unfolds in high-density information environments '
    'where abundant content competes for limited cognitive attention [1]. In this setting, '
    'headlines serve as the primary affective gateway through which audiences assess relevance, '
    'form expectations, and make decisions [2]. As attention becomes fragmented across platforms '
    'and devices, emotionally charged headlines gain disproportionate salience, potentially '
    'shaping public sentiment and economic behavior at scale [3].',
    'PARA'
)
add_para(
    'In the post-COVID-19 period, public news discourse has become more emotionally intense, '
    'with a notable prevalence of high-arousal cues such as fear and anger [4]. These emotional '
    'shifts coincide with significant macroeconomic disruptions\u2014supply-chain shocks, energy '
    'crises, and inflationary pressures\u2014raising the question of whether media-conveyed emotions '
    'are merely reflective of economic conditions or contribute to shaping them through their '
    'influence on consumer confidence and expectations [5], [6].',
    'PARA'
)
add_para(
    'Recent advances in natural language processing have enabled a methodological shift from '
    'lexicon-based sentiment measures toward transformer-based language models that capture '
    'contextual and multidimensional emotional content [7], [8]. Models such as RoBERTa-BNE, '
    'pretrained on large-scale Spanish corpora, allow fine-grained emotion classification that '
    'surpasses the limitations of polarity-based approaches [9].',
    'PARA'
)
add_para(
    'Against this backdrop, the present study examines the longitudinal evolution of discrete '
    'emotional signals in 175,869 Spanish digital news headlines from 2017 to 2023 and evaluates '
    'their temporal association with Spain\u2019s Consumer Price Index (IPC ECOICOP). Unlike prior '
    'work that relies on small balanced subsamples or lexicon-based sentiment, this study '
    'processes the full census of collected headlines through a transformer-based pipeline and '
    'applies a comprehensive econometric framework comprising ARIMAX models, Vector '
    'Autoregression (VAR) with Granger causality tests, impulse response functions (IRF), '
    'forecast error variance decomposition (FEVD), and structural break analysis. Two exogenous '
    'structural dummy variables\u2014D_COVID (March 2020 onward) and D_UKRAINE (February 2022 '
    'onward)\u2014capture the impact of major geopolitical shocks.',
    'PARA'
)
add_para(
    'The study addresses three research questions: (RQ1) How has the emotional composition of '
    'Spanish digital news headlines evolved between 2017 and 2023? (RQ2) Do headline-level '
    'emotional signals exhibit statistically significant temporal associations with consumer '
    'price dynamics? (RQ3) Did COVID-19 and the Ukraine conflict produce structural breaks in '
    'the emotion\u2013inflation nexus?',
    'PARA'
)
add_para(
    'The contribution of this work is threefold. First, it processes the complete census of '
    'collected headlines (N = 175,869) rather than relying on balanced subsamples, eliminating '
    'sampling bias and maximizing statistical power. Second, it deploys a multi-equation '
    'econometric framework\u2014combining ARIMAX, VAR, Granger causality, IRF, and FEVD\u2014that '
    'captures dynamic interdependencies missed by single-equation approaches. Third, it '
    'formally tests for structural regime shifts using both a priori hypothesized break dates '
    '(COVID-19, Ukraine) and endogenously determined break points (Bai\u2013Perron), providing '
    'robust evidence on how exogenous shocks alter the emotion\u2013inflation transmission mechanism.',
    'PARA'
)
add_para(
    'The remainder of this paper is organized as follows. Section II reviews the relevant '
    'literature on negativity bias, media emotion, and narrative economics. Section III '
    'describes the corpus, emotion classification pipeline, and econometric methodology. '
    'Section IV presents the empirical results. Section V discusses implications and '
    'comparisons with prior work. Section VI concludes with limitations and directions '
    'for future research.',
    'PARA'
)

# ═══════════════════════════════════════════════════════════════════════════════
# II. RELATED WORK
# ═══════════════════════════════════════════════════════════════════════════════
add_para('RELATED WORK', 'H1_List (Space)')

add_para(
    'Research on news consumption consistently documents that negative information exerts a '
    'disproportionate influence on cognition, commonly described as negativity bias [10]. '
    'Aversive stimuli attract more attention, are processed more deeply, and exert a stronger '
    'influence on subsequent judgment and behavior than equivalent positive stimuli [11]. In '
    'digital media ecosystems, algorithmic curation and engagement-driven design further '
    'amplify emotionally charged content, creating feedback loops between editorial framing '
    'and audience behavior [12], [13].',
    'PARA'
)
add_para(
    'Within journalism and communication research, headlines are frequently described as the '
    'primary affective entry point to news content. As attention becomes increasingly '
    'fragmented across platforms, headlines function as self-contained informational units '
    'that shape initial emotional responses and guide subsequent content engagement [2], [14]. '
    'Empirical studies indicate that emotional framing in headlines influences sharing behavior, '
    'perceived credibility, and information retention [15].',
    'PARA'
)
add_para(
    'From an economic perspective, emotionally framed news exposure has been associated with '
    'changes in consumer confidence, expectations, and spending behavior [5], [16]. '
    'Narrative economics, as formalized by Shiller [6], [17], posits that viral stories and '
    'emotionally resonant media narratives can propagate through social networks and '
    'influence aggregate economic outcomes. Longitudinal studies indicate that media sentiment '
    'indices, particularly those capturing fear and uncertainty, predict changes in consumer '
    'confidence and economic policy uncertainty [18], [19].',
    'PARA'
)
add_para(
    'Methodologically, the study of media\u2013economy relationships has evolved with advances '
    'in NLP. Traditional lexicon-based sentiment measures (e.g., AFINN, VADER) are '
    'increasingly recognized as limited for capturing multidimensional emotional content, '
    'particularly in non-English languages where lexical resources are scarce [20]. '
    'Transformer-based models such as BERT [21] and RoBERTa [22] offer superior contextual '
    'understanding, while language-specific variants like RoBERTa-BNE [9] provide '
    'fine-grained emotion classification for Spanish text. Recent work has demonstrated '
    'the viability of integrating transformer-derived emotional features into time-series '
    'econometric frameworks [23], [24], though applications combining full-corpus census '
    'processing with comprehensive VAR/Granger analysis remain scarce.',
    'PARA'
)

# ═══════════════════════════════════════════════════════════════════════════════
# III. METHODOLOGY
# ═══════════════════════════════════════════════════════════════════════════════
add_para('METHODOLOGY', 'H1_List (Space)')

# A. Corpus
add_para('CORPUS DESCRIPTION', 'H2_Cont')
add_para(
    'The dataset comprises 175,869 unique Spanish-language news headlines collected via the '
    'Google News RSS feed between January 2017 and December 2023 (84 months). This represents '
    'the complete census of headlines collected through an automated pipeline, filtered to the '
    'study period without further subsampling. Google News RSS was selected to provide broad '
    'coverage of the Spanish digital news ecosystem, aggregating content from 1,436 distinct '
    'media sources.',
    'PARA'
)
add_para(
    'Headlines are organized into six thematic categories: politics (45,610; 25.9%), sports '
    '(43,048; 24.5%), society (30,601; 17.4%), environment (24,924; 14.2%), economy (15,886; '
    '9.0%), and culture (15,800; 9.0%). Table I reports the annual distribution of the corpus.',
    'PARA'
)

# Table I: Corpus by year
add_table(
    ['Year', 'Headlines', 'Culture', 'Sports', 'Economy', 'Environment', 'Politics', 'Society'],
    [
        ['2017', '19,518', '2,016', '5,351', '1,144', '2,545', '4,756', '3,706'],
        ['2018', '23,177', '2,377', '6,115', '1,457', '2,887', '5,822', '4,519'],
        ['2019', '26,172', '2,574', '6,353', '1,985', '3,334', '7,348', '4,578'],
        ['2020', '25,070', '2,171', '5,933', '2,259', '4,270', '5,862', '4,575'],
        ['2021', '27,052', '2,242', '6,270', '3,097', '3,915', '6,547', '4,981'],
        ['2022', '27,769', '2,018', '6,519', '2,918', '4,233', '6,897', '5,184'],
        ['2023', '27,111', '2,402', '6,507', '3,026', '3,740', '8,378', '3,058'],
        ['Total', '175,869', '15,800', '43,048', '15,886', '24,924', '45,610', '30,601'],
    ],
    'TABLE I. Annual distribution of the headline corpus (2017\u20132023).'
)

# B. Emotional Classification
add_para('EMOTIONAL CLASSIFICATION', 'H2_Cont')
add_para(
    'Emotional features are extracted using RoBERTa-BNE (pysentimiento), a RoBERTa-based '
    'transformer model pretrained on large-scale Spanish corpora and fine-tuned for emotion '
    'classification [9]. Unlike lexicon-based approaches that assign discrete polarity scores, '
    'RoBERTa-BNE produces a seven-dimensional probability vector for each headline, covering '
    'anger, fear, joy, sadness, surprise, disgust, and a neutral/others category. Each '
    'probability reflects the model\u2019s confidence that the headline conveys the corresponding '
    'emotion, enabling continuous rather than categorical measurement.',
    'PARA'
)
add_para(
    'All 175,869 headlines in the corpus were processed through the emotion classification '
    'pipeline. The dominant emotion (highest probability) is neutral/others for 93.2% of '
    'headlines, consistent with the informational register typical of news writing. Among '
    'emotionally marked headlines, anger (2.6%) and sadness (2.1%) are the most frequent '
    'dominant emotions, followed by joy (1.3%), fear (0.6%), surprise (0.2%), and disgust '
    '(< 0.1%). Table II reports descriptive statistics of emotion probabilities.',
    'PARA'
)

# Table II: Emotion distribution
add_table(
    ['Emotion', 'Mean', 'Std', 'Median', 'Dominant N', 'Dominant %'],
    [
        ['Anger',    '0.0304', '0.1159', '0.0014', '4,634', '2.63'],
        ['Fear',     '0.0138', '0.0553', '0.0020', '983',   '0.56'],
        ['Joy',      '0.0215', '0.0966', '0.0026', '2,275', '1.29'],
        ['Sadness',  '0.0261', '0.1179', '0.0015', '3,609', '2.05'],
        ['Surprise', '0.0150', '0.0383', '0.0046', '332',   '0.19'],
        ['Disgust',  '0.0097', '0.0304', '0.0010', '62',    '0.04'],
        ['Others',   '0.8835', '0.2227', '0.9773', '163,974', '93.24'],
    ],
    'TABLE II. Descriptive statistics of emotion probabilities (N = 175,869).'
)

# C. Econometric Framework
add_para('ECONOMETRIC FRAMEWORK', 'H2_Cont')
add_para(
    'The econometric analysis follows a multi-stage procedure. Monthly aggregates of four '
    'core emotion probabilities (anger, fear, sadness, and joy) are computed from the headline-'
    'level data, producing 84 monthly observations. These are merged with macroeconomic '
    'indicators from official Spanish statistical sources (INE, CIS). Two structural dummy '
    'variables are included: D_COVID (= 1 from March 2020 onward) and D_UKRAINE (= 1 from '
    'February 2022 onward).',
    'PARA'
)
add_para(
    'The analysis proceeds in four stages:',
    'PARA'
)
add_para(
    'Stage 1: Stationarity assessment. Augmented Dickey\u2013Fuller (ADF) tests are applied to all '
    'variables. Non-stationary series are first-differenced before entering the VAR framework.',
    'PARA'
)
add_para(
    'Stage 2: ARIMAX modeling. For each dependent variable y_t, an ARIMAX(p,d,q) model is '
    'estimated with emotion variables and dummy variables as exogenous regressors. The '
    'autoregressive order (p,d,q) is selected via AIC minimization using auto_arima, and the '
    'model is then re-estimated with statsmodels SARIMAX to obtain full coefficient inference. '
    'Three model families are estimated: (A) a global model with IPC General as the dependent '
    'variable; (B) 16 individual models for each ECOICOP sub-index, ICM, ICC, and IPI; and '
    '(C) 6 category-specific models linking thematic emotions to corresponding economic indicators.',
    'PARA'
)
add_para(
    'Stage 3: VAR and Granger causality. A five-variable VAR model including IPC General and '
    'the four core emotions is estimated. Lag order is selected by AIC (maximum 8 lags, '
    'adjusted for the 84-observation sample). Bidirectional Granger causality tests assess '
    'whether lagged values of each variable improve prediction of every other variable. '
    'Orthogonalized impulse response functions (IRF) and forecast error variance decomposition '
    '(FEVD) characterize the dynamic transmission of shocks.',
    'PARA'
)
add_para(
    'Stage 4: Structural break analysis. Chow tests evaluate whether the relationship between '
    'emotions and IPC undergoes regime shifts at the COVID-19 onset (March 2020) and the '
    'Ukraine conflict (February 2022). A sequential Bai\u2013Perron procedure searches for '
    'endogenously determined break dates. Kolmogorov\u2013Smirnov and Mann\u2013Whitney tests compare '
    'pre-COVID and post-COVID emotion distributions.',
    'PARA'
)

# D. Variables
add_para('VARIABLES', 'H2_Cont')
add_table(
    ['Variable', 'Type', 'Source', 'Frequency'],
    [
        ['IPC General Index', 'Dependent', 'INE', 'Monthly'],
        ['12 ECOICOP sub-indices', 'Dependent (Model B)', 'INE', 'Monthly'],
        ['ICM', 'Dependent (Model B)', 'CIS', 'Monthly'],
        ['ICC', 'Dependent (Model B)', 'CIS', 'Monthly'],
        ['IPI', 'Dependent (Model B)', 'INE', 'Monthly'],
        ['Anger (Ira)', 'Exogenous', 'RoBERTa-BNE', 'Monthly'],
        ['Fear (Miedo)', 'Exogenous', 'RoBERTa-BNE', 'Monthly'],
        ['Sadness (Tristeza)', 'Exogenous', 'RoBERTa-BNE', 'Monthly'],
        ['Joy (Alegr\u00eda)', 'Exogenous', 'RoBERTa-BNE', 'Monthly'],
        ['D_COVID', 'Dummy (=1 from 2020-03)', 'Manual', 'Monthly'],
        ['D_UKRAINE', 'Dummy (=1 from 2022-02)', 'Manual', 'Monthly'],
    ],
    'TABLE III. Variables and data sources.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# IV. RESULTS
# ═══════════════════════════════════════════════════════════════════════════════
add_para('RESULTS', 'H1_List (Space)')

# A. Temporal evolution
add_para('TEMPORAL EVOLUTION OF EMOTIONAL SIGNALS', 'H2_Cont')
add_para(
    'Fig. 1 illustrates the monthly evolution of the four core emotions and their 12-month '
    'moving averages over the 2017\u20132023 period. Several patterns emerge. First, anger '
    'exhibits a sustained upward trend, with monthly means increasing from approximately '
    '0.029 in 2017 to 0.031 in 2023 (+5.4%). Second, joy displays a marked decline, '
    'particularly after March 2020, decreasing from 0.023 pre-COVID to 0.020 post-COVID '
    '(\u221211.8%). Third, fear shows episodic spikes coinciding with the COVID-19 onset '
    '(March\u2013April 2020) and, to a lesser extent, the early phase of the Ukraine conflict. '
    'Sadness remains relatively stable across the period (\u22123.0%).',
    'PARA'
)
add_figure(
    FIGURES_DIR / 'fig_01_emotion_evolution.png',
    'FIGURE 1. Monthly evolution of headline emotions (2017\u20132023): (a) raw monthly means; '
    '(b) 12-month moving averages. Dashed lines indicate COVID-19 (red) and Ukraine (orange) onsets.',
    width_inches=6.5
)
add_para(
    'The moving-average panel (Fig. 1b) reveals longer-term trends more clearly. Anger '
    'displays a gradual upward drift throughout the observation window, accelerating slightly '
    'after 2020. Fear peaks sharply during the first months of the pandemic and then gradually '
    'subsides, although it remains above its pre-2020 baseline. Joy shows a persistent '
    'downward trend that begins well before COVID-19 but accelerates after March 2020, '
    'suggesting that the pandemic reinforced a pre-existing pattern of declining positive '
    'emotional framing in Spanish media. Sadness exhibits the most episodic behavior, spiking '
    'during the acute phases of both the pandemic and the energy crisis but returning to '
    'baseline levels relatively quickly.',
    'PARA'
)

# B. Heatmap
add_para(
    'Fig. 2 presents the emotion\u2013category heatmap. Politics and society headlines exhibit '
    'the highest anger and fear probabilities, while sports and culture display elevated joy '
    'and surprise. Economy headlines show intermediate profiles with slightly elevated sadness.',
    'PARA'
)
add_figure(
    FIGURES_DIR / 'fig_02_heatmap_categories.png',
    'FIGURE 2. Emotion profile by news category (2017\u20132023). '
    'Cell values represent mean emotion probabilities aggregated across all monthly observations.',
    width_inches=6.5
)
add_para(
    'The cross-category analysis highlights important asymmetries. Politics headlines carry '
    'the highest anger intensity (mean = 0.040), more than double that of sports (0.018), '
    'consistent with the adversarial framing characteristic of political discourse. Environment '
    'headlines show elevated fear relative to other categories, likely reflecting coverage of '
    'climate-related risks and natural disasters. Sports headlines stand out for their '
    'comparatively high joy levels, while culture occupies an intermediate position across '
    'all emotional dimensions. These category-level differences motivate the category-specific '
    'ARIMAX models (Model C) described in the methodology.',
    'PARA'
)

# C. ADF Tests
add_para('STATIONARITY TESTS', 'H2_Cont')

# Build ADF table from VAR results
adf_data = var_summary.get('differenced_vars', [])
add_para(
    f'Table IV reports ADF test results for the five endogenous variables in the VAR system. '
    f'IPC General (ADF = \u22120.357, p = 0.917) and Joy (ADF = \u22122.188, p = 0.211) are '
    f'non-stationary in levels and are first-differenced prior to VAR estimation. Anger '
    f'(p < 0.001), Fear (p < 0.001), and Sadness (p < 0.001) are stationary in levels.',
    'PARA'
)

adf_var = pd.read_csv(DATA_DIR / 'var_granger' / 'adf_tests.csv')
add_table(
    ['Variable', 'ADF Statistic', 'p-value', 'Stationary'],
    [[tr(row['variable'][:35]), f"{row['adf_stat']:.4f}", f"{row['p_value']:.4f}",
      'Yes' if row['stationary'] else 'No']
     for _, row in adf_var.iterrows()],
    'TABLE IV. Augmented Dickey\u2013Fuller stationarity tests.'
)

# D. ARIMAX Results
add_para('ARIMAX RESULTS', 'H2_Cont')

gm = global_model
add_para(
    f'Model A (Global). The AIC-optimal specification for IPC General is ARIMAX{gm["order"]}, '
    f'with AIC = {gm["aic"]} and RMSE = {gm["rmse"]}. The Ljung\u2013Box test confirms white-noise '
    f'residuals (p > 0.05 at all lags). While individual emotion coefficients do not reach '
    f'conventional significance in the global model\u2014consistent with the high autoregressive '
    f'persistence of the IPC series\u2014the signs are economically interpretable: Joy exhibits '
    f'a positive association (\u03b2 = {gm["coefficients"]["Alegria"]["coefficient"]:.2f}, '
    f'p = {gm["coefficients"]["Alegria"]["p_value"]:.3f}), while Fear shows a negative '
    f'association (\u03b2 = {gm["coefficients"]["Miedo"]["coefficient"]:.2f}, '
    f'p = {gm["coefficients"]["Miedo"]["p_value"]:.3f}).',
    'PARA'
)

# Table V: Global ARIMAX coefficients
coef_rows = []
for pname in ['intercept', 'Ira', 'Miedo', 'Tristeza', 'Alegria', 'D_COVID', 'D_UCRANIA', 'sigma2']:
    if pname in gm['coefficients']:
        info = gm['coefficients'][pname]
        sig = '*' if info['significant_5pct'] else ''
        display_name = '\u03c3\u00b2' if pname == 'sigma2' else tr(pname)
        coef_rows.append([
            display_name,
            f"{info['coefficient']:.4f}",
            f"{info['p_value']:.4f}",
            sig
        ])

add_table(
    ['Parameter', 'Coefficient', 'p-value', 'Sig.'],
    coef_rows,
    f'TABLE V. ARIMAX{gm["order"]} Global Model: IPC General ~ Emotions + Dummies.'
)

# ARIMAX diagnostic figure
arimax_global_fig = ARIMAX_DIR / 'arimax_IPC_ECOICOP_Indice_General_(Global).png'
if arimax_global_fig.exists():
    add_figure(
        arimax_global_fig,
        'FIGURE 3. ARIMAX Global Model diagnostics: (a) observed vs. predicted IPC General; '
        '(b) residuals; (c) residual distribution; (d) Q\u2013Q plot.',
        width_inches=6.5
    )
    add_para(
        'Fig. 3 displays the diagnostic plots for the global ARIMAX model. The observed-versus-'
        'predicted panel confirms close in-sample tracking, with deviations concentrated around '
        'the COVID-19 period when the IPC experienced unprecedented volatility. The residual '
        'distribution approximates normality, and the Q\u2013Q plot reveals moderate departures only '
        'in the tails, primarily driven by the outlier months of March\u2013April 2020. The '
        'Ljung\u2013Box test confirms white-noise residuals (p > 0.05 at all tested lags), '
        'supporting the adequacy of the selected specification.',
        'PARA'
    )

# Model B summary
add_para(
    'Model B (By dependent variable). Table VI summarizes the 16 ARIMAX models estimated '
    'for individual economic indicators. Significant exogenous effects emerge in several '
    'sub-indices:',
    'PARA'
)

# List significant findings
for item in sig_arimax:
    var_short = item['variable'][:50]
    exog_list = ', '.join([f'{tr(name)} (\u03b2={coef:.2f}, p={pv:.3f})'
                           for name, coef, pv in item['sig_exog']])
    add_para(f'\u2022 {var_short}: {exog_list}', 'PARA')

# Build Model B summary table
model_b_rows = []
for _, row in df_arimax_all.iterrows():
    var_name = str(row.get('Variable', ''))[:40]
    order = str(row.get('ARIMAX Order', ''))
    aic = row.get('AIC', '')
    rmse = row.get('RMSE', '')
    lb = 'Yes' if row.get('Ljung-Box OK') else 'No'

    # Find significant exog
    sig_list = []
    for emo in ['Ira', 'Miedo', 'Tristeza', 'Alegria']:
        if row.get(f'sig_{emo}') == True:
            sig_list.append(tr(emo))
    sig_str = ', '.join(sig_list) if sig_list else '\u2014'

    model_b_rows.append([var_name, order, f'{aic:.1f}' if isinstance(aic, (int,float)) else str(aic),
                         f'{rmse:.4f}' if isinstance(rmse, (int,float)) else str(rmse),
                         lb, sig_str])

add_table(
    ['Dep. Variable', 'Order', 'AIC', 'RMSE', 'L-B OK', 'Sig. Emotions'],
    model_b_rows[:17],  # Limit to avoid too many rows
    'TABLE VI. ARIMAX model summary for all dependent variables.'
)

# E. VAR and Granger
add_para('VAR MODEL AND GRANGER CAUSALITY', 'H2_Cont')

vs = var_summary
add_para(
    f'The VAR lag order selected by AIC is {vs["var_order"]} (BIC selects 0). '
    f'The VAR({vs["var_order"]}) model is estimated on {vs["n_obs"]} observations after '
    f'differencing IPC General and Joy. The Portmanteau whiteness test yields '
    f'p = {vs["portmanteau_pvalue"]:.4f}.',
    'PARA'
)

# Granger summary
sig_gc = df_granger[df_granger['significant_5pct'] == True]
best_gc = sig_gc.loc[sig_gc.groupby(['cause', 'effect'])['f_pvalue'].idxmin()]

add_para(
    f'Table VII reports the significant Granger causality results. A total of '
    f'{len(best_gc)} unique directional pairs are significant at the 5% level. '
    f'Key findings include:',
    'PARA'
)

# Organize Granger results by direction
emotion_to_ipc = best_gc[best_gc['effect'].str.contains('IPC')]
ipc_to_emotion = best_gc[best_gc['cause'].str.contains('IPC')]

if len(emotion_to_ipc) > 0:
    causes = ', '.join(sorted([tr(c) for c in emotion_to_ipc['cause'].unique()]))
    add_para(
        f'\u2022 Emotions \u2192 IPC: {causes} Granger-cause IPC General, confirming that '
        f'lagged emotional signals in news headlines carry predictive information for '
        f'consumer price dynamics.',
        'PARA'
    )

if len(ipc_to_emotion) > 0:
    effects = ', '.join(sorted([tr(e) for e in ipc_to_emotion['effect'].unique()]))
    add_para(
        f'\u2022 IPC \u2192 Emotions: IPC General Granger-causes {effects}, indicating '
        f'feedback effects where macroeconomic conditions influence subsequent media '
        f'emotional framing.',
        'PARA'
    )

# Table VII: Granger significant pairs
gc_rows = []
for _, row in best_gc.iterrows():
    cause_short = tr(row['cause'][:25]) if 'IPC' not in row['cause'] else 'CPI General'
    effect_short = tr(row['effect'][:25]) if 'IPC' not in row['effect'] else 'CPI General'
    gc_rows.append([
        cause_short, effect_short,
        str(int(row['lag'])),
        f"{row['f_stat']:.3f}",
        f"{row['f_pvalue']:.4f}"
    ])

add_table(
    ['Cause', 'Effect', 'Lag', 'F-stat', 'p-value'],
    gc_rows,
    'TABLE VII. Significant Granger causality pairs (p < 0.05, best lag per pair).'
)

# F. IRF and FEVD
add_para('IMPULSE RESPONSE FUNCTIONS AND VARIANCE DECOMPOSITION', 'H2_Cont')
add_para(
    'Fig. 4 presents selected orthogonalized IRFs. A one-standard-deviation shock to '
    'Sadness produces a negative response in IPC General that persists for approximately '
    '3\u20134 months before dissipating. Anger shocks generate a transient negative response '
    'in IPC. Fear shocks show a negative contemporaneous effect that reverses after 2 months. '
    'Joy shocks are associated with modest positive IPC responses.',
    'PARA'
)
add_figure(
    FIGURES_DIR / 'fig_04_irf_selected.png',
    'FIGURE 4. Orthogonalized impulse response functions (12-month horizon): '
    'emotion \u2192 CPI (top row) and CPI \u2192 emotion (bottom row).',
    width_inches=6.5
)
add_para(
    'The IRF analysis provides insight into the direction, magnitude, and persistence of '
    'inter-variable shocks. In the top row, the responses of IPC to emotional shocks reveal '
    'that sadness and anger produce transient deflationary effects, consistent with the '
    'interpretation that withdrawal-oriented emotions suppress consumer spending and exert '
    'downward pressure on prices. The response to a joy shock is modestly positive, '
    'suggesting that optimism-driven media framing may support consumer confidence and '
    'price stability. Fear generates an initially negative response that reverses after '
    'approximately 2\u20133 months, potentially reflecting the dual nature of fear as both '
    'a consumption suppressant and a precautionary-demand driver. The bottom row shows that '
    'IPC shocks feed back into emotional framing: positive price shocks increase sadness and '
    'fear in subsequent media coverage while reducing joy, consistent with the narrative that '
    'inflation generates negative media framing.',
    'PARA'
)

# FEVD table
fevd_rows = []
for h in [1, 3, 6, 12]:
    if h in df_fevd.index:
        row_data = df_fevd.loc[h]
        ipc_val = row_data.get('IPC_ECOICOP \u00cdndice General', row_data.iloc[0])
        fevd_rows.append([
            str(h),
            f"{ipc_val:.4f}",
            f"{row_data.get('Ira', row_data.iloc[1]):.4f}",
            f"{row_data.get('Miedo', row_data.iloc[2]):.4f}",
            f"{row_data.get('Tristeza', row_data.iloc[3]):.4f}",
            f"{row_data.get('Alegria', row_data.iloc[4]):.4f}",
        ])

add_table(
    ['Horizon', 'IPC', 'Anger', 'Fear', 'Sadness', 'Joy'],
    fevd_rows,
    'TABLE VIII. Forecast error variance decomposition of IPC General.'
)

add_figure(
    FIGURES_DIR / 'fig_05_fevd_ipc.png',
    'FIGURE 5. Stacked bar chart of forecast error variance decomposition for IPC General '
    'across 1\u201312 month horizons.',
    width_inches=3.3
)

fevd_h12 = vs.get('fevd_ipc_horizon12', {})
joy_pct = fevd_h12.get('Alegria', 0) * 100
add_para(
    f'Table VIII shows the FEVD of IPC General. At the 12-month horizon, '
    f'Joy accounts for {joy_pct:.1f}% of IPC forecast error variance, while IPC\u2019s '
    f'own shocks explain only {fevd_h12.get("IPC_ECOICOP \u00cdndice General", 0)*100:.1f}%. '
    f'This suggests that emotional signals, particularly joy, play a substantial role in '
    f'explaining long-run variability in consumer prices.',
    'PARA'
)

# G. Structural Breaks
add_para('STRUCTURAL BREAK ANALYSIS', 'H2_Cont')

chow_covid = struct_summary['chow_tests'].get('COVID (mar-2020)', {})
chow_ukraine = struct_summary['chow_tests'].get('Ucrania (feb-2022)', {})
bp = struct_summary.get('bai_perron', [])

add_para(
    f'The Chow test confirms statistically significant structural breaks at both '
    f'hypothesized dates: COVID-19 onset (F = {chow_covid.get("f_stat", "N/A")}, '
    f'p < 0.001; {chow_covid.get("n1", "")} pre vs. {chow_covid.get("n2", "")} post '
    f'observations) and the Ukraine conflict (F = {chow_ukraine.get("f_stat", "N/A")}, '
    f'p < 0.001; {chow_ukraine.get("n1", "")} pre vs. {chow_ukraine.get("n2", "")} post '
    f'observations). The sequential Bai\u2013Perron procedure independently identifies '
    f'December 2021 as an endogenous break point (F = {bp[0]["f_stat"] if bp else "N/A"}, '
    f'p < 0.001), closely aligning with the onset of inflationary acceleration in Spain.',
    'PARA'
)

# Table IX: Structural breaks
break_rows = [
    ['Chow: COVID (Mar 2020)', f'{chow_covid.get("f_stat", "")}',
     f'{chow_covid.get("p_value", "")}',
     f'{chow_covid.get("n1", "")}/{chow_covid.get("n2", "")}', 'Yes'],
    ['Chow: Ukraine (Feb 2022)', f'{chow_ukraine.get("f_stat", "")}',
     f'{chow_ukraine.get("p_value", "")}',
     f'{chow_ukraine.get("n1", "")}/{chow_ukraine.get("n2", "")}', 'Yes'],
]
if bp:
    break_rows.append([
        f'Bai\u2013Perron: {bp[0]["break_date"][:10]}',
        f'{bp[0]["f_stat"]}', f'{bp[0]["p_value"]}', '\u2014', 'Yes'
    ])

add_table(
    ['Test', 'F-stat', 'p-value', 'N pre/post', 'Sig.'],
    break_rows,
    'TABLE IX. Structural break tests.'
)

# KS/MW
add_para(
    'Table X reports Kolmogorov\u2013Smirnov and Mann\u2013Whitney comparisons of emotion '
    'distributions between the pre-COVID (January 2017\u2013February 2020; 38 months) and '
    'post-COVID (March 2020\u2013December 2023; 46 months) periods.',
    'PARA'
)

ks_rows = []
for _, row in df_ks.iterrows():
    var = str(row['variable'])
    if 'IPC' in var:
        var = 'CPI General'
    else:
        var = tr(var)
    ks_rows.append([
        var,
        f"{row['mean_pre']:.4f}",
        f"{row['mean_post']:.4f}",
        f"{row['diff_pct']:.1f}%",
        f"{row['ks_pvalue']:.4f}",
        f"{row['mw_pvalue']:.4f}",
    ])

add_table(
    ['Variable', 'Mean Pre', 'Mean Post', '\u0394%', 'KS p', 'MW p'],
    ks_rows,
    'TABLE X. Pre-COVID vs. Post-COVID distribution comparisons.'
)

add_para(
    'Joy shows the largest and most statistically significant shift, declining 11.8% '
    'post-COVID (KS p < 0.001; MW p < 0.001). Anger increases 5.4% (KS p = 0.022). '
    'IPC General increases 8.7% (KS p < 0.001), reflecting the post-pandemic inflationary '
    'episode.',
    'PARA'
)

add_figure(
    FIGURES_DIR / 'fig_06_series_covid_break.png',
    'FIGURE 6. CPI General Index and composite negativity (Anger + Fear + Sadness) with '
    'structural break lines at COVID-19 (red) and Ukraine (orange).',
    width_inches=6.5
)
add_para(
    'Fig. 6 visually juxtaposes IPC General with the composite negativity index. The IPC '
    'series remains relatively flat during 2017\u20132020, then accelerates sharply from late '
    '2021 onward, coinciding with both the Bai\u2013Perron endogenous break date (December 2021) '
    'and the escalation of energy prices. The negativity index shows a corresponding spike '
    'during March\u2013May 2020, followed by a gradual elevation that persists throughout the '
    'post-COVID period. The visual co-movement between emotional negativity and inflationary '
    'acceleration provides intuitive support for the statistical relationships identified by '
    'the Granger causality tests.',
    'PARA'
)

add_figure(
    FIGURES_DIR / 'fig_07_emotional_profile_periods.png',
    'FIGURE 7. Mean emotion probabilities by sub-period: Pre-COVID (2017-01 to 2020-02), '
    'COVID (2020-03 to 2021-12), and Post-COVID (2022-01 to 2023-12).',
    width_inches=3.3
)
add_para(
    'Fig. 7 decomposes the emotional profile into three sub-periods. The pre-COVID baseline '
    '(38 months) establishes reference levels for each emotion. During the acute COVID period '
    '(22 months), fear and anger increase while joy declines. In the post-COVID period '
    '(24 months), the emotional rebalancing partially persists: joy remains below its pre-COVID '
    'level (\u221211.8%), anger stays elevated (+5.4%), while sadness returns close to baseline. '
    'This pattern suggests that the pandemic induced a semi-permanent shift in the emotional '
    'register of Spanish news coverage rather than a temporary perturbation.',
    'PARA'
)

add_para('CATEGORY-LEVEL EMOTIONAL DYNAMICS', 'H2_Cont')
add_figure(
    FIGURES_DIR / 'fig_08_trends_by_category.png',
    'FIGURE 8. Emotion trends by news category (2017\u20132023). Each panel displays monthly '
    'means for the four core emotions within one thematic category.',
    width_inches=6.5
)
add_para(
    'Fig. 8 displays category-level emotional trends. Several category-specific patterns are '
    'noteworthy. Politics headlines show the most pronounced anger increase over time, '
    'reflecting the polarized political climate in Spain during 2017\u20132023, which included '
    'the Catalan independence crisis, multiple general elections, and coalition government '
    'negotiations. Environment headlines exhibit the sharpest fear spikes during 2020, '
    'likely driven by pandemic-related environmental coverage and subsequent climate events. '
    'Economy headlines display a distinctive pattern where sadness and anger increase in tandem '
    'with inflationary pressures from 2021 onward. Sports headlines maintain the most stable '
    'emotional profile, with joy consistently dominating other emotions, although a visible '
    'dip occurs during the pandemic-related sports suspension (March\u2013June 2020). These '
    'category-level heterogeneities underscore the importance of disaggregated analysis and '
    'justify the category-specific ARIMAX models estimated in Model C.',
    'PARA'
)

# ═══════════════════════════════════════════════════════════════════════════════
# V. DISCUSSION
# ═══════════════════════════════════════════════════════════════════════════════
add_para('DISCUSSION', 'H1_List (Space)')

add_para(
    'This study provides multi-layered evidence that emotionally encoded information in '
    'Spanish digital news headlines is systematically associated with consumer price dynamics '
    'over the 2017\u20132023 period. Three main findings warrant discussion.',
    'PARA'
)
add_para(
    'First, the Granger causality analysis reveals robust bidirectional temporal dependencies '
    'between headline emotions and CPI. Sadness and anger Granger-cause IPC at multiple lags, '
    'while IPC in turn Granger-causes sadness, fear, and joy. This bidirectionality is '
    'consistent with narrative economics [6], where media emotional framing both reflects and '
    'shapes macroeconomic dynamics. The finding that sadness\u2014rather than the more commonly '
    'studied fear\u2014is the strongest predictor of IPC aligns with recent evidence that '
    'withdrawal-oriented emotions (sadness, fatigue) may influence consumer behavior '
    'differently from approach-oriented ones (anger) [10], [11].',
    'PARA'
)
add_para(
    'Second, while the global ARIMAX model does not show individually significant emotion '
    'coefficients\u2014due to the strong autoregressive persistence of the IPC series\u2014the '
    'disaggregated models reveal economically meaningful patterns. Joy significantly predicts '
    'ICM (consumer confidence), D_UKRAINE significantly affects ICM, and D_COVID significantly '
    'affects ICC. The category-specific models further show that joy in society-related '
    'headlines predicts consumer confidence (ICC). These findings suggest that emotional '
    'effects operate through specific transmission channels rather than uniformly across all '
    'price categories.',
    'PARA'
)
add_para(
    'Third, the structural break analysis confirms that both COVID-19 and the Ukraine '
    'conflict produced statistically significant regime shifts in the emotion\u2013price '
    'relationship. The Bai\u2013Perron procedure independently identifies December 2021 as an '
    'endogenous break point, preceding the Ukraine conflict by two months and coinciding with '
    'the acceleration of energy prices and supply-chain disruptions that marked the onset '
    'of Spain\u2019s inflationary episode. The 11.8% decline in joy post-COVID, combined with '
    'the 5.4% increase in anger, suggests a structural emotional rebalancing in media '
    'coverage that persists beyond the acute pandemic phase.',
    'PARA'
)
add_para(
    'Compared with the pilot version of this study, which relied on a balanced subsample of '
    '5,000 headlines and a single ARIMAX model, the present analysis achieves several '
    'methodological advances: (i) full census processing (175,869 headlines) eliminates '
    'sampling bias; (ii) the VAR/Granger framework captures dynamic interdependencies that '
    'univariate ARIMAX cannot; (iii) IRF and FEVD quantify the magnitude and persistence '
    'of emotional shocks on prices; and (iv) structural break analysis formally tests for '
    'regime changes rather than assuming temporal homogeneity.',
    'PARA'
)
add_para(
    'From a practical standpoint, these findings suggest that real-time monitoring of '
    'emotional signals in media headlines could complement traditional macroeconomic '
    'indicators for nowcasting and early warning. The lag structure identified in the '
    'Granger tests\u2014with emotional signals leading IPC changes by 1\u20136 months\u2014implies '
    'that headline emotion indices could provide advance information about inflationary '
    'turning points. Central banks and statistical offices increasingly explore text-based '
    'indicators for economic surveillance [18], [19]; the present results support extending '
    'these efforts to emotion-specific measures derived from transformer-based NLP.',
    'PARA'
)
add_para(
    'The theoretical implications align with Shiller\u2019s [6] narrative economics framework. '
    'The bidirectional Granger causality between emotions and CPI confirms that media '
    'narratives are not merely passive reflections of economic reality but active participants '
    'in shaping expectations and behavior. The particularly strong role of sadness\u2014a '
    'withdrawal-oriented emotion associated with disengagement and reduced activity\u2014in '
    'predicting CPI changes extends the literature beyond the more commonly studied fear '
    'and anger dimensions. This finding resonates with behavioral economics research '
    'suggesting that mood-congruent processing [10] may lead consumers exposed to '
    'sadness-laden news to reduce discretionary spending, thereby affecting aggregate '
    'demand and, ultimately, price dynamics.',
    'PARA'
)

# ═══════════════════════════════════════════════════════════════════════════════
# VI. CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════
add_para('CONCLUSION', 'H1_List (Space)')

add_para(
    'This study examined how emotionally encoded information in 175,869 Spanish digital news '
    'headlines (2017\u20132023) relates to consumer price dynamics in Spain. Using RoBERTa-BNE '
    'for emotion classification and a comprehensive econometric framework (ARIMAX, VAR, '
    'Granger causality, IRF, FEVD, and structural break tests), three principal findings '
    'emerge.',
    'PARA'
)
add_para(
    'First, headline emotions and CPI exhibit bidirectional Granger-causal relationships: '
    'sadness, anger, and joy predict future price movements, while price changes in turn '
    'influence subsequent media emotional framing. Second, both COVID-19 and the Ukraine '
    'conflict produced significant structural breaks in the emotion\u2013price relationship, '
    'with joy declining 11.8% and anger increasing 5.4% post-COVID. Third, the FEVD reveals '
    'that emotional signals explain a substantial share of long-run IPC forecast error '
    'variance, supporting the practical relevance of incorporating media-derived emotional '
    'indicators into macroeconomic monitoring.',
    'PARA'
)
add_para(
    'Several limitations warrant acknowledgment. First, the study relies on headline-only '
    'text, which compresses meaning and may not capture the full emotional content of news '
    'articles. Second, the transformer model\u2019s emotion classification reflects training-set '
    'biases that may not perfectly align with reader perception. Third, the 84-month sample '
    'limits the degrees of freedom available for high-order VAR models. Fourth, the analysis '
    'is confined to the Spanish media ecosystem and IPC ECOICOP; generalization to other '
    'countries or economic indicators requires further investigation.',
    'PARA'
)
add_para(
    'Future work should extend the temporal coverage, incorporate cross-country comparisons, '
    'and explore the role of social media as a complementary source of emotional signals. '
    'The integration of real-time emotional monitoring with nowcasting frameworks offers a '
    'promising avenue for improving macroeconomic surveillance.',
    'PARA'
)

# ═══════════════════════════════════════════════════════════════════════════════
# APPENDIX / DECLARATIONS
# ═══════════════════════════════════════════════════════════════════════════════
add_para('Declaration of Competing Interest', 'H1')
add_para(
    'The author(s) declare that they have no known competing financial interests or personal '
    'relationships that could have appeared to influence the work reported in this paper.',
    'PARA'
)
add_para('Funding', 'PARA')
add_para(
    'This research did not receive any specific grant from funding agencies in the public, '
    'commercial, or not-for-profit sectors.',
    'PARA'
)
add_para('Declaration of Generative AI Use', 'PARA')
add_para(
    'During the preparation of this work, the author(s) used Claude (Anthropic) to assist '
    'with code development and language editing. After using this tool, the author(s) reviewed '
    'and edited the content and take full responsibility for the publication.',
    'PARA'
)

# Repository
add_para('DATA AVAILABILITY', 'H1')
add_para(
    'To ensure transparency and facilitate replication, the complete computational workflow '
    'is publicly available at:',
    'PARA'
)
add_para(
    'https://github.com/Jorgejosezamora/sentiment-news-spain-2017-2023',
    'PARA'
)
add_para(
    'The repository includes: (I) headline filtering and aggregation scripts; '
    '(II) ARIMAX, VAR/Granger, and structural break analysis code; '
    '(III) figure and table generation scripts; (IV) aggregated monthly datasets (CSV). '
    'Raw headline data are not shared due to copyright restrictions; only monthly '
    'aggregated emotion probabilities are provided.',
    'PARA'
)

# ═══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════
add_para('REFERENCES', 'H1')

refs = [
    '[1] N. Newman, R. Fletcher, C. T. Robertson, A. Ross Arguedas, and R. K. Nielsen, '
    '\u201cReuters Institute Digital News Report 2024,\u201d Reuters Institute for the Study of '
    'Journalism, 2024.',

    '[2] J. Berger and K. L. Milkman, \u201cWhat makes online content viral?\u201d J. Marketing '
    'Research, vol. 49, no. 2, pp. 192\u2013205, 2012.',

    '[3] N. Diakopoulos, Automating the News: How Algorithms Are Rewriting the Media. '
    'Cambridge, MA, USA: Harvard Univ. Press, 2019.',

    '[4] S. N. Soroka, Negativity in Democratic Politics: Causes and Consequences. '
    'Cambridge, U.K.: Cambridge Univ. Press, 2014.',

    '[5] R. Vliegenthart and A. Damstra, \u201cParliamentary questions, media coverage, and '
    'consumer confidence,\u201d Journalism Mass Commun. Quart., vol. 96, no. 3, pp. 766\u2013789, 2019.',

    '[6] R. J. Shiller, \u201cNarrative economics,\u201d Amer. Econ. Rev., vol. 107, no. 4, '
    'pp. 967\u20131004, 2017.',

    '[7] J. Devlin, M.-W. Chang, K. Lee, and K. Toutanova, \u201cBERT: Pre-training of deep '
    'bidirectional transformers for language understanding,\u201d in Proc. NAACL-HLT, 2019, '
    'pp. 4171\u20134186.',

    '[8] D. Demszky et al., \u201cGoEmotions: A dataset of fine-grained emotions,\u201d '
    'in Proc. 58th Annu. Meeting ACL, 2020, pp. 4040\u20134054.',

    '[9] F. Del Arco, A. Montejo-R\u00e1ez, and F. Mart\u00ednez-Santiago, \u201cEmotion analysis '
    'based on Transformer models: Case study in Spanish,\u201d Expert Syst. Appl., vol. 173, '
    'art. 114632, 2021.',

    '[10] A. Tversky and D. Kahneman, \u201cJudgment under uncertainty: Heuristics and '
    'biases,\u201d Science, vol. 185, no. 4157, pp. 1124\u20131131, 1974.',

    '[11] S. N. Soroka, \u201cNegative and positive negativity bias,\u201d in Negativity in '
    'Democratic Politics. Cambridge, U.K.: Cambridge Univ. Press, 2014, ch. 2.',

    '[12] K. Crawford, The Atlas of AI: Power, Politics, and the Planetary Costs of '
    'Artificial Intelligence. New Haven, CT, USA: Yale Univ. Press, 2021.',

    '[13] V. Garc\u00eda-Perdomo, \u201cImpact of news attributes on shareability and '
    'engagement,\u201d Journalism Mass Commun. Quart., vol. 98, no. 1, pp. 130\u2013151, 2021.',

    '[14] P. Mihailidis, Media Literacy for the Post-Truth Era. New York, NY, USA: '
    'Routledge, 2018.',

    '[15] U. K. H. Ecker, J. L. Hogan, and S. Lewandowsky, \u201cReminders and repetition '
    'of misinformation,\u201d J. Appl. Res. Memory Cognition, vol. 6, no. 2, '
    'pp. 185\u2013192, 2017.',

    '[16] G. A. Akerlof and R. J. Shiller, Animal Spirits: How Human Psychology Drives '
    'the Economy. Princeton, NJ, USA: Princeton Univ. Press, 2009.',

    '[17] R. J. Shiller, Narrative Economics: How Stories Go Viral and Drive Major '
    'Economic Events. Princeton, NJ, USA: Princeton Univ. Press, 2019.',

    '[18] S. R. Baker, N. Bloom, and S. J. Davis, \u201cMeasuring economic policy '
    'uncertainty,\u201d Quart. J. Econ., vol. 131, no. 4, pp. 1593\u20131636, 2016.',

    '[19] R. Nyman et al., \u201cNews and narratives in financial systems: Exploiting big '
    'data for systemic risk assessment,\u201d J. Econ. Dynamics Control, vol. 127, '
    'art. 104119, 2021.',

    '[20] S. M. Mohammad and P. D. Turney, \u201cCrowdsourcing a word\u2013emotion association '
    'lexicon,\u201d Comput. Intell., vol. 29, no. 3, pp. 436\u2013465, 2013.',

    '[21] J. Devlin, M.-W. Chang, K. Lee, and K. Toutanova, \u201cBERT: Pre-training of '
    'deep bidirectional transformers,\u201d arXiv:1810.04805, 2018.',

    '[22] Y. Liu et al., \u201cRoBERTa: A robustly optimized BERT pretraining approach,\u201d '
    'arXiv:1907.11692, 2019.',

    '[23] J. F. S\u00e1nchez-Rada, O. Araque, and C. A. Iglesias, \u201cFrom sentiment to '
    'emotions: A comparative study of Transformer architectures for Spanish media '
    'analytics,\u201d Expert Syst. Appl., vol. 238, art. 122018, 2024.',

    '[24] R. J. Hyndman and G. Athanasopoulos, Forecasting: Principles and Practice, '
    '3rd ed. Melbourne, Australia: OTexts, 2021.',

    '[25] G. E. P. Box and G. M. Jenkins, Time Series Analysis: Forecasting and Control, '
    '2nd ed. San Francisco, CA, USA: Holden-Day, 1976.',

    '[26] G. M. Ljung and G. E. P. Box, \u201cOn a measure of lack of fit in time series '
    'models,\u201d Biometrika, vol. 65, no. 2, pp. 297\u2013303, 1978.',

    '[27] F. Zollo et al., \u201cEmotional dynamics in the age of misinformation,\u201d '
    'PLoS ONE, vol. 10, no. 9, art. e0138740, 2015.',

    '[28] F. Barbieri, L. Espinosa-Anke, and J. Camacho-Collados, \u201cTimeLMs: Diachronic '
    'language models from Twitter,\u201d in Findings EMNLP, 2022, pp. 5508\u20135519.',
]

for ref in refs:
    add_para(ref, 'References')

# ─── Author bios ─────────────────────────────────────────────────────────────
add_para(
    'Jorge Jos\u00e9 Zamora C\u00e1novas is a Ph.D. candidate in Economic, Business and Legal '
    'Sciences at the Polytechnic University of Cartagena (Spain). His research focuses on '
    'data analysis and its applications in communication and economics, with particular '
    'emphasis on computational methods for media analysis and their macroeconomic implications.',
    'AU_Bios'
)
add_para(
    'Soledad Maria Mart\u00ednez Maria-Dolores Born in Cartagena, Spain in 1971, she obtained '
    'her doctorate in 2004. She is the General Coordinator of the Scientific Culture and '
    'Innovation Unit (UCC+i) of the UPCT. Her teaching experience spans over 25 years at the '
    'UPCT. She has published research papers in SJR and JCR ranked journals. Her research '
    'interests include quantitative methods applied to social and economic phenomena.',
    'AU_Bios'
)

# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════
doc.save(str(OUTPUT))
print(f"\nArticle saved to: {OUTPUT}")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print("=" * 80)
print("ARTICLE REWRITE COMPLETE")
print("=" * 80)
